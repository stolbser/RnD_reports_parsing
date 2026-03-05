import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import json
import os


class TitleSheetGenerator:
    def __init__(self, root):
        self.root = root
        self.root.title("Генератор титульных листов НИР")
        self.root.geometry("900x750")

        # Загрузка шаблонов
        self.load_templates()
        self.create_widgets()

    def load_templates(self):
        """Загрузка шаблонов из файла"""
        self.templates_file = "templates.json"

        default_templates = {
            "organizations": [
                "ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ «ИССЛЕДОВАТЕЛЬСКИЙ ЦЕНТР ТМК» (ООО «ИЦ ТМК»)"
            ],
            "udk_list": ["669-153.4"],
            "approvers": [
                {
                    "name": "С.И. Благовещенский",
                    "position": "И.о Главного инженера Филиала ПАО «ТМК» ВТЗ",
                    "label": "Согласовал"
                },
                {
                    "name": "И.Ю. Пышминцев",
                    "position": "Генеральный директор ООО «ИЦ ТМК»",
                    "degree": "д-р техн. наук",
                    "label": "УТВЕРЖДАЮ"
                }
            ],
            "topics": [
                "Исследование материалов (сталей, сплавов и защитных покрытий), применяемых при производстве новых видов трубной продукции и разработка новых технологий и оборудования (промежуточный)"
            ],
            "leaders": [
                {
                    "name": "А.Н. Мальцева",
                    "position": "Руководитель НИР, начальник отдела перспективных и функциональных",
                    "degree": "канд. техн. наук"
                }
            ],
            "contract_numbers": ["11/23ИЦ"],
            "performers": [
                {
                    "name": "А.М. Арсенкин",
                    "position": "Заведующий лабораторией металловедения и прочности",
                    "degree": "канд. техн. наук"
                },
                {
                    "name": "С.Д. Столбов",
                    "position": "Научный сотрудник лаборатории металловедения и прочности",
                    "degree": "канд. техн. наук"
                },
                {
                    "name": "А.В. Дмитриев",
                    "position": "Инженер-исследователь лаборатории металловедения и прочности",
                    "degree": ""
                },
                {
                    "name": "И.А. Темченко",
                    "position": "Техник-лаборант лаборатории металловедения и прочности",
                    "degree": ""
                },
                {
                    "name": "А.А. Французов",
                    "position": "Нормоконтроль",
                    "degree": ""
                }
            ]
        }

        if os.path.exists(self.templates_file):
            with open(self.templates_file, 'r', encoding='utf-8') as f:
                self.templates = json.load(f)

            # Добавляем недостающие ключи
            updated = False
            for key, value in default_templates.items():
                if key not in self.templates:
                    self.templates[key] = value
                    updated = True

            if updated:
                self.save_templates()
        else:
            self.templates = default_templates
            self.save_templates()

    def save_templates(self):
        """Сохранение шаблонов в файл"""
        with open(self.templates_file, 'w', encoding='utf-8') as f:
            json.dump(self.templates, f, ensure_ascii=False, indent=2)

    def create_widgets(self):
        """Создание элементов интерфейса"""
        main_frame = ttk.Frame(self.root, padding="10")
        main_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))

        # Настройка растягивания
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        main_frame.columnconfigure(1, weight=1)

        # Организация
        ttk.Label(main_frame, text="Организация:", font=('Arial', 10, 'bold')).grid(row=0, column=0, sticky=tk.W, pady=5)
        self.org_combo = ttk.Combobox(main_frame, width=70, values=self.templates["organizations"], font=('Arial', 9))
        self.org_combo.grid(row=0, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        if self.templates["organizations"]:
            self.org_combo.current(0)

        # К
        ttk.Label(main_frame, text="УДК:", font=('Arial', 10, 'bold')).grid(row=1, column=0, sticky=tk.W, pady=5)
        self.udk_combo = ttk.Combobox(main_frame, width=30, values=self.templates["udk_list"], font=('Arial', 9))
        self.udk_combo.grid(row=1, column=1, sticky=tk.W, pady=5)
        if self.templates["udk_list"]:
            self.udk_combo.current(0)

        # Согласующий
        ttk.Label(main_frame, text="Согласовал:", font=('Arial', 10, 'bold')).grid(row=2, column=0, sticky=tk.W, pady=5)
        self.approver1_combo = ttk.Combobox(main_frame, width=70, font=('Arial', 9))
        self.approver1_combo.grid(row=2, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.populate_approver_combo(self.approver1_combo, "Согласовал")

        # Утверждающий
        ttk.Label(main_frame, text="Утвердил:", font=('Arial', 10, 'bold')).grid(row=3, column=0, sticky=tk.W, pady=5)
        self.approver2_combo = ttk.Combobox(main_frame, width=70, font=('Arial', 9))
        self.approver2_combo.grid(row=3, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.populate_approver_combo(self.approver2_combo, "УТВЕРЖДАЮ")

        # Номер договора
        ttk.Label(main_frame, text="Номер договора:", font=('Arial', 10, 'bold')).grid(row=4, column=0, sticky=tk.W,
                                                                                       pady=5)
        self.contract_combo = ttk.Combobox(main_frame, width=30, values=self.templates["contract_numbers"],
                                           font=('Arial', 9))
        self.contract_combo.grid(row=4, column=1, sticky=tk.W, pady=5)
        if self.templates["contract_numbers"]:
            self.contract_combo.current(0)

        # Дополнительное соглашение
        ttk.Label(main_frame, text="Доп. соглашение №:", font=('Arial', 10, 'bold')).grid(row=5, column=0, sticky=tk.W,
                                                                                          pady=5)
        self.addendum_num_entry = ttk.Entry(main_frame, width=15, font=('Arial', 9))
        self.addendum_num_entry.grid(row=5, column=1, sticky=tk.W, pady=5)
        self.addendum_num_entry.insert(0, "38")

        ttk.Label(main_frame, text="от:", font=('Arial', 10, 'bold')).grid(row=5, column=1, sticky=tk.E, pady=5)
        self.addendum_date_entry = ttk.Entry(main_frame, width=15, font=('Arial', 9))
        self.addendum_date_entry.grid(row=5, column=2, sticky=tk.W, pady=5)
        self.addendum_date_entry.insert(0, "17.02.2026")

        # Тема НИР
        ttk.Label(main_frame, text="Тема НИР:", font=('Arial', 10, 'bold')).grid(row=6, column=0, sticky=tk.NW, pady=5)
        self.topic_combo = ttk.Combobox(main_frame, width=70, values=self.templates["topics"], font=('Arial', 9))
        self.topic_combo.grid(row=6, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        if self.templates["topics"]:
            self.topic_combo.current(0)

        # Руководитель работ
        ttk.Label(main_frame, text="Руководитель работ:", font=('Arial', 10, 'bold')).grid(row=7, column=0, sticky=tk.W,
                                                                                           pady=5)
        self.leader_combo = ttk.Combobox(main_frame, width=70, font=('Arial', 9))
        self.leader_combo.grid(row=7, column=1, columnspan=2, sticky=(tk.W, tk.E), pady=5)
        self.populate_leader_combo()

        # Год
        ttk.Label(main_frame, text="Год:", font=('Arial', 10, 'bold')).grid(row=8, column=0, sticky=tk.W, pady=5)
        self.year_entry = ttk.Entry(main_frame, width=10, font=('Arial', 9))
        self.year_entry.grid(row=8, column=1, sticky=tk.W, pady=5)
        self.year_entry.insert(0, "2026")

        # Список исполнителей
        ttk.Label(main_frame, text="Исполнители:", font=('Arial', 10, 'bold')).grid(row=9, column=0, sticky=tk.NW, pady=10)

        performers_frame = ttk.Frame(main_frame)
        performers_frame.grid(row=9, column=1, columnspan=2, sticky=(tk.W, tk.E, tk.N, tk.S), pady=5)

        self.performers_listbox = tk.Listbox(performers_frame, width=70, height=10, selectmode=tk.MULTIPLE, font=('Arial', 9))
        self.performers_listbox.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(performers_frame, orient=tk.VERTICAL, command=self.performers_listbox.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.performers_listbox.config(yscrollcommand=scrollbar.set)

        for perf in self.templates["performers"]:
            display_text = f"{perf['position']}" + (f", {perf['degree']}" if perf.get('degree') else "") + f" - {perf['name']}"
            self.performers_listbox.insert(tk.END, display_text)
            self.performers_listbox.select_set(0, tk.END)  # <--- Выделить всех по умолчанию

        # Кнопки управления
        btn_frame = ttk.Frame(main_frame)
        btn_frame.grid(row=10, column=0, columnspan=3, pady=20)

        ttk.Button(btn_frame, text="📄 Создать титульный лист", command=self.generate_document, width=25).pack(
            side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="➕ Добавить шаблон", command=self.open_template_editor, width=20).pack(side=tk.LEFT, padx=5)
        ttk.Button(btn_frame, text="❌ Выход", command=self.root.quit, width=15).pack(side=tk.LEFT, padx=5)

    def populate_approver_combo(self, combo, label_filter):
        """Заполнение комбобокса согласующих/утверждающих"""
        values = []
        for appr in self.templates["approvers"]:
            if appr.get("label") == label_filter:
                display = f"{appr['position']}" + (f", {appr['degree']}" if appr.get('degree') else "") + f" - {appr['name']}"
                values.append(display)
        combo['values'] = values
        if values:
            combo.current(0)

    def populate_leader_combo(self):
        """Заполнение комбобокса руководителей"""
        values = []
        for lead in self.templates["leaders"]:
            display = f"{lead['position']}" + (f", {lead['degree']}" if lead['degree'] else "") + f" - {lead['name']}"
            values.append(display)
        self.leader_combo['values'] = values
        if values:
            self.leader_combo.current(0)

    def remove_cell_borders(self, table):
        """Удаление границ таблицы"""
        tbl = table._tbl
        for cell in tbl.iter_tcs():
            tcPr = cell.tcPr
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                cell.append(tcPr)

            tcBorders = OxmlElement('w:tcBorders')

            for border_name in ['top', 'left', 'bottom', 'right']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tcBorders.append(border)

            tcPr.append(tcBorders)

    def generate_document(self):
        """Генерация документа Word"""
        try:
            # Получение данных из формы
            organization = self.org_combo.get()
            udk = self.udk_combo.get()

            # СТАЛО (правильно)
            # Получаем текст из выпадающего меню
            approver1_text = self.approver1_combo.get()
            approver1 = None
            for appr in self.templates["approvers"]:
                if appr.get("label") == "Согласовал":
                    # Формируем строку так же, как в populate_approver_combo
                    display = f"{appr['position']}" + (
                        f", {appr['degree']}" if appr.get('degree') else "") + f" - {appr['name']}"
                    if display == approver1_text:
                        approver1 = appr
                        break

            # Получаем текст из выпадающего меню
            approver2_text = self.approver2_combo.get()
            approver2 = None
            for appr in self.templates["approvers"]:
                if appr.get("label") == "УТВЕРЖДАЮ":
                    # Формируем строку так же, как в populate_approver_combo
                    display = f"{appr['position']}" + (
                        f", {appr['degree']}" if appr.get('degree') else "") + f" - {appr['name']}"
                    if display == approver2_text:
                        approver2 = appr
                        break

            contract_num = self.contract_combo.get()
            addendum_num = self.addendum_num_entry.get()
            addendum_date = self.addendum_date_entry.get()
            topic = self.topic_combo.get()

            leader_idx = self.leader_combo.current()
            leader = self.templates["leaders"][leader_idx] if leader_idx >= 0 else None

            year = self.year_entry.get()

            # Выбранные исполнители
            selected_indices = self.performers_listbox.curselection()
            performers = []
            for i in selected_indices:
                try:
                    idx = int(i)  # <--- Преобразуем строку в число
                    # Проверяем, что индекс в пределах массива
                    if idx < len(self.templates["performers"]):
                        performers.append(self.templates["performers"][idx])
                except IndexError:
                    # Если индекс выходит за границы, пропускаем
                    continue

            # Создание документа
            doc = Document()

            # Настройка полей
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(3)
                section.right_margin = Cm(1.5)

            # Создание титульного листа
            self.create_title_page(doc, organization, udk, approver1, approver2, contract_num, addendum_num, addendum_date, topic, leader, year)

            # Создание списка исполнителей
            self.create_performers_list(doc, performers)

            # Сохранение документа
            filename = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word документ", "*.docx")],
                initialfile=f"Титульный_лист_{year}.docx",
                title="Сохранить титульный лист"
            )

            if filename:
                doc.save(filename)
                messagebox.showinfo("✅ Успех", f"Документ успешно создан:\n{filename}")

        except Exception as e:
            messagebox.showerror("❌ Ошибка", f"Ошибка при создании документа:\n{str(e)}")

    def create_title_page(self, doc, organization, udk, approver1, approver2,
                          contract_num, addendum_num, addendum_date, topic, leader, year):
        """Создание титульного листа ТОЧЬ-В-ТОЧЬ как в образце"""

        # 1. Организация (1 раз)
        p = doc.add_paragraph()
        p.alignment = 1

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run(organization)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False

        # 3. Пустая строка
        doc.add_paragraph()

        # 2. УДК (1 раз)
        p = doc.add_paragraph()
        p.alignment = 0

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run(f"УДК {udk}")
        run.font.name = 'Times New Roman'
        run.size = Pt(12)
        run.font.bold = False

        # 3. Пустая строка
        doc.add_paragraph()


        # 4. Таблица 1 строка × 2 колонки (без границ)
        table = doc.add_table(rows=1, cols=2)
        table.autofit = False

        # Установка ширины колонок
        col_widths = [Cm(10), Cm(9)]
        for i, width in enumerate(col_widths):
            for cell in table.columns[i].cells:
                cell.width = width

        # Удаление границ таблицы
        self.remove_cell_borders(table)

        # Колонка 1: СОГЛАСОВАЛ
        cell1 = table.cell(0, 0)
        p1 = cell1.paragraphs[0]
        p1.alignment = 0

        # Убираем интервалы
        p1.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p1.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p1.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run1a = p1.add_run("СОГЛАСОВАНО\n")
        run1a.font.name = 'Times New Roman'
        run1a.font.size = Pt(12)
        run1a.font.bold = False

        run1b = p1.add_run(f"{approver1['position']}\n")
        run1b.font.name = 'Times New Roman'
        run1b.font.size = Pt(12)
        run1b.font.bold = False

        run1c = p1.add_run("_" * 15)
        run1c.font.name = 'Times New Roman'
        run1c.font.size = Pt(12)
        run1c.font.bold = False

        run1d = p1.add_run(f"{approver1['name']}\n")
        run1d.font.name = 'Times New Roman'
        run1d.font.size = Pt(12)
        run1d.font.bold = False

        run1e = p1.add_run(f"«____» _____________ {year} г.")
        run1e.font.name = 'Times New Roman'
        run1e.font.size = Pt(12)
        run1e.font.bold = False

        # Колонка 2: УТВЕРЖДАЮ
        cell2 = table.cell(0, 1)
        p2 = cell2.paragraphs[0]
        p2.alignment = 0

        # Убираем интервалы
        p2.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p2.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p2.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run2a = p2.add_run("УТВЕРЖДАЮ\n\n")
        run2a.font.name = 'Times New Roman'
        run2a.font.size = Pt(12)
        run2a.font.bold = False

        pos_text = approver2['position']
        if approver2.get('degree'):
            pos_text += f" {approver2['degree']}"

        run2b = p2.add_run(pos_text + "\n")
        run2b.font.name = 'Times New Roman'
        run2b.font.size = Pt(12)
        run2b.font.bold = False

        run2c = p2.add_run("_" * 15)
        run2c.font.name = 'Times New Roman'
        run2c.font.size = Pt(12)
        run2c.font.bold = False

        run2d = p2.add_run(f"{approver2['name']}\n")
        run2d.font.name = 'Times New Roman'
        run2d.font.size = Pt(12)
        run2d.font.bold = False

        run2e = p2.add_run(f"«____» __________ {year} г.")
        run2e.font.name = 'Times New Roman'
        run2e.font.size = Pt(12)
        run2e.font.bold = False

        # 5. Пустая строка
        doc.add_paragraph()
        doc.add_paragraph()

        # 6. Название отчета (1 раз)
        p = doc.add_paragraph()
        p.alignment = 1

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run("ОТЧЕТ\nО НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False

        p = doc.add_paragraph()
        p.alignment = 1

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run("по теме:")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False

        p = doc.add_paragraph()

        p.alignment = 1

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run(topic)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False

        p = doc.add_paragraph()
        p.alignment = 1

        # Убираем интервалы
        p.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run = p.add_run(f"(договор № {contract_num}, Дополнительное соглашение №{addendum_num} от {addendum_date} г.)")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = False

        # 7. Пустая строка
        doc.add_paragraph()
        doc.add_paragraph()

        # 8. Таблица руководителя НИР (1 строка × 2 колонки)
        leader_table = doc.add_table(rows=1, cols=2)
        leader_table.autofit = False

        # Установка ширины колонок
        leader_col_widths = [Cm(9), Cm(9)]
        for i, width in enumerate(leader_col_widths):
            for cell in leader_table.columns[i].cells:
                cell.width = width

        # Удаление границ
        self.remove_cell_borders(leader_table)

        # Левая колонка - должность
        cell_leader_left = leader_table.cell(0, 0)
        p_leader_left = cell_leader_left.paragraphs[0]
        p_leader_left.alignment = 0

        pos_text = leader['position']
        if leader['degree']:
            pos_text += f", {leader['degree']}"

        run_leader_pos = p_leader_left.add_run(pos_text)
        run_leader_pos.font.name = 'Times New Roman'
        run_leader_pos.font.size = Pt(12)
        run_leader_pos.font.bold = False

        # Правая колонка - подпись и фамилия
        cell_leader_right = leader_table.cell(0, 1)
        p_leader_right = cell_leader_right.paragraphs[0]
        p_leader_right.alignment = 1

        run_leader_line = p_leader_right.add_run("_" * 30)
        run_leader_line.font.name = 'Times New Roman'
        run_leader_line.font.size = Pt(12)
        run_leader_line.font.bold = False

        p_leader_right.add_run("\n\n")

        run_leader_name = p_leader_right.add_run(leader['name'])
        run_leader_name.font.name = 'Times New Roman'
        run_leader_name.font.size = Pt(12)
        run_leader_name.font.bold = False

        # 9. Пустая строка
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        # 10. Город и год (по центру)
        p_city = doc.add_paragraph()
        p_city.alignment = 1

        # Убираем интервалы
        p_city.paragraph_format.space_after = Pt(0)  # Интервал после абзаца
        p_city.paragraph_format.space_before = Pt(0)  # Интервал перед абзацем
        p_city.paragraph_format.line_spacing = 1  # Межстрочный интервал (1 = одинарный)

        run_city = p_city.add_run(f"Москва {year}")
        run_city.font.name = 'Times New Roman'
        run_city.font.size = Pt(12)
        run_city.font.bold = False

        # 11. Разрыв страницы
        doc.add_page_break()


    def create_performers_list(self, doc, performers):
        print(f"Начинаю генерацию списка исполнителей. Количество исполнителей: {len(performers)}")
        for i, perf in enumerate(performers):
            print(f"Исполнитель {i}: {perf['name']}")

        """Создание списка исполнителей"""
        # Заголовок
        p = doc.add_paragraph()
        p.alignment = 1
        run = p.add_run("СПИСОК ИСПОЛНИТЕЛЕЙ")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True

        doc.add_paragraph()

        # Таблица исполнителей (БЕЗ ГРАНИЦ!)
        if performers:
            table = doc.add_table(rows=len(performers), cols=3)
            table.autofit = False

            # Установка ширины колонок
            col_widths = [Cm(9), Cm(4), Cm(5)]
            for i, width in enumerate(col_widths):
                for cell in table.columns[i].cells:
                    cell.width = width

            # Удаление границ
            self.remove_cell_borders(table)

            for i, perf in enumerate(performers):
                # Должность
                cell0 = table.cell(i, 0)
                p0 = cell0.paragraphs[0]
                p0.paragraph_format.space_after = Pt(30)
                p0.alignment = 0

                pos_text = perf['position']
                if perf.get('degree'):
                    pos_text += f", {perf['degree']}"

                run0 = p0.add_run(pos_text)
                run0.font.name = 'Times New Roman'
                run0.font.size = Pt(12)
                run0.font.bold = False

                # Подпись и дата
                cell1 = table.cell(i, 1)
                p1 = cell1.paragraphs[0]
                p1.paragraph_format.space_after = Pt(30)
                p1.alignment = 1

                run1 = p1.add_run("_" * 17 + "\nподпись, дата")
                run1.font.name = 'Times New Roman'
                run1.font.size = Pt(12)
                run1.font.bold = False

                # ФИО
                cell2 = table.cell(i, 2)
                p2 = cell2.paragraphs[0]
                p2.paragraph_format.space_after = Pt(30)
                p2.alignment = 0

                run2 = p2.add_run(perf['name'])
                run2.font.name = 'Times New Roman'
                run2.font.size = Pt(12)
                run2.font.bold = False


    def open_template_editor(self):
        """Открытие окна редактора шаблонов"""
        editor_window = tk.Toplevel(self.root)
        editor_window.title("Редактор шаблонов")
        editor_window.geometry("800x600")

        notebook = ttk.Notebook(editor_window)
        notebook.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Вкладка Организации
        org_frame = ttk.Frame(notebook)
        notebook.add(org_frame, text="Организации")

        ttk.Label(org_frame, text="Добавить организацию:").pack(pady=5)
        org_entry = ttk.Entry(org_frame, width=70)
        org_entry.pack(pady=5)

        def add_organization():
            org = org_entry.get().strip()
            if org and org not in self.templates["organizations"]:
                self.templates["organizations"].append(org)
                self.save_templates()
                self.org_combo['values'] = self.templates["organizations"]
                messagebox.showinfo("✅", "Организация добавлена!")
                org_entry.delete(0, tk.END)
            elif org in self["organizations"]:
                messagebox.showwarning("⚠️", "Такая организация уже есть!")
            else:
                messagebox.showwarning("⚠️", "Введите название организации!")

        ttk.Button(org_frame, text="Добавить", command=add_organization).pack(pady=5)

        # Список организаций
        ttk.Label(org_frame, text="Существующие организации:").packady = 10
        org_listbox = tk.Listbox(org_frame, width=80, height=10)
        org_listbox.pack(fill=tk.BOTH, expand=True, pady=5)

        for org in self.templates["organizations"]:
            org_listbox.insert(tk.END, org)

        # Вкладка Подписанты
        approver_frame = ttk.Frame(notebook)
        notebook.add(approver_frame, text="Подписанты")

        ttk.Label(approver_frame, text="ФИО:").grid(row=0, column=0, sticky=tk.W, pady=5)
        name_entry = ttk.Entry(approver_frame, width=50)
        name_entry.grid(row=0, column=1, pady=5)

        ttk.Label(approver_frame, text="Должность:").grid(row=1, column=0, sticky=tk.W, pady=5)
        position_entry = ttk.Entry(approver_frame, width=50)
        position_entry.grid(row=1, column=1, pady=5)

        ttk.Label(approver_frame, text="Ученая степень:").grid(row=2, column=0, sticky=tk.W, pady=5)
        degree_entry = ttk.Entry(approver_frame, width=50)
        degree_entry.grid(row=2, column=1, pady=5)

        ttk.Label(approver_frame, text="Тип:").grid(row=3, column=0, sticky=tk.W, pady=5)
        label_combo = ttk.Combobox(approver_frame, values=["Согласовал", "УТВЕРЖДАЮ"], width=47)
        label_combo.grid(row=3, column=1, pady=5)
        label_combo.current(1)

        def add_approver():
            name = name_entry.get().strip()
            position = position_entry.get().strip()
            degree = degree_entry.get().strip()
            label = label_combo.get()

            if not name or not position or not label:
                messagebox.showwarning("⚠️", "Заполните все обязательные поля!")
                return

            new_approver = {"name": name, "position": position, "label": label}
            if degree:
                new_approver["degree"] = degree

            self.templates["approvers"].append(new_approver)
            self.save_templates()
            self.populate_approver_combo(self.approver1_combo, "Согласовал")
            self.populate_approver_combo(self.approver2_combo, "УТВЕРЖДАЮ")
            messagebox.showinfo("✅", "Подписант добавлен!")

            name_entry.delete(0, tk.END)
            position_entry.delete(0, tk.END)
            degree_entry.delete(0, tk.END)

        ttk.Button(approver_frame, text="Добавить подписанта", command=add_approver).grid(row=4, column=0, columnspan=2,
                                                                                          pady=10)

        # Вкладка Темы
        topic_frame = ttk.Frame(notebook)
        notebook.add(topic_frame, text="Темы НИР")

        ttk.Label(topic_frame, text="Добавить тему:").pack(pady=5)
        topic_entry = ttk.Entry(topic_frame, width=70)
        topic_entry.pack(pady=5)

        def add_topic():
            topic = topic_entry.get().strip()
            if topic and topic not in self.templates["topics"]:
                self.templates["topics"].append(topic)
                self.save_templates()
                self.topic_combo['values'] = self.templates["topics"]
                messagebox.showinfo("✅", "Тема добавлена!")
                topic_entry.delete(0, tk.END)
            elif topic in self.templates["topics"]:
                messagebox.showwarning("⚠️", "Такая тема уже есть!")
            else:
                messagebox.showwarning("⚠️", "Введите тему!")

        ttk.Button(topic_frame, text="Добавить", command=add_topic).pack(pady=5)

        # Вкладка Руководители
        leader_frame = ttk.Frame(notebook)
        notebook.add(leader_frame, text="Руководители")

        ttk.Label(leader_frame, text="ФИО:").grid(row=0, column=0, sticky=tk.W, pady=5)
        lname_entry = ttk.Entry(leader_frame, width=50)
        lname_entry.grid(row=0, column=1, pady=5)

        ttk.Label(leader_frame, text="Должность:").grid(row=1, column=0, sticky=tk.W, pady=5)
        lposition_entry = ttk.Entry(leader_frame, width=50)
        lposition_entry.grid(row=1, column=1, pady=5)

        ttk.Label(leader_frame, text="Ученая степень:").grid(row=2, column=0, sticky=tk.W, pady=5)
        ldegree_entry = ttk.Entry(leader_frame, width=50)
        ldegree_entry.grid(row=2, column=1, pady=5)

        def add_leader():
            name = lname_entry.get().strip()
            position = lposition_entry.get().strip()
            degree = ldegree_entry.get().strip()

            if not name or not position:
                messagebox.showwarning("⚠️", "Заполните ФИО и должность!")
                return

            new_leader = {"name": name, "position": position, "degree": degree}
            self.templates["leaders"].append(new_leader)
            self.save_templates()
            self.populate_leader_combo()
            messagebox.showinfo("✅", "Руководитель добавлен!")

            lname_entry.delete(0, tk.END)
            lposition_entry.delete(0, tk.END)
            ldegree_entry.delete(0, tk.END)

        ttk.Button(leader_frame, text="Добавить руководителя", command=add_leader).grid(row=3, column=0, columnspan=2, pady=10)

        # Вкладка Исполнители
        performer_frame = ttk.Frame(notebook)
        notebook.add(performer_frame, text="Исполнители")

        ttk.Label(performer_frame, text="ФИО:").grid(row=0, column=0, sticky=tk.W, pady=5)
        pname_entry = ttk.Entry(performer_frame, width=50)
        pname_entry.grid(row=0, column=1, pady=5)

        ttk.Label(performer_frame, text="Должность:").grid(row=1, column=0, sticky=tk.W, pady=5)
        pposition_entry = ttk.Entry(performer_frame, width=50)
        pposition_entry.grid(row=1, column=1, pady=5)

        ttk.Label(performer_frame, text="Ученая степень:").grid(row=2, column=0, sticky=tk.W, pady=5)
        pdegree_entry = ttk.Entry(performer_frame, width=50)
        pdegree_entry.grid(row=2, column=1, pady=5)

        def add_performer():
            name = pname_entry.get().strip()
            position = pposition_entry.get().strip()
            degree = pdegree_entry.get().strip()

            if not name or not position:
                messagebox.showwarning("⚠️", "Заполните ФИО и должность!")
                return

            new_performer = {"name": name, "position": position, "degree": degree}
            self.templates["performers"].append(new_performer)
            self.save_templates()

            # Обновление списка в основном окне
            self.performers_listbox.delete(0, tk.END)
            for perf in self.templates["performers"]:
                display_text = f"{perf['position']}" + (f", {perf['degree']}" if perf['degree'] else "") + f" - {perf['name']}"
                self.performers_listbox.insert(tk.END, display_text)

            messagebox.showinfo("✅", "Исполнитель добавлен!")

            pname_entry.delete(0, tk.END)
            pposition_entry.delete(0, tk.END)
            pdegree_entry.delete(0, tk.END)

        ttk.Button(performer_frame, text="Добавить исполнителя", command=add_performer).grid(row=3, column=0, columnspan=2, pady=10)

        ttk.Button(editor_window, text="Закрыть", command=editor_window.destroy).pack(pady=10)

if __name__ == "__main__":
    root = tk.Tk()
    app = TitleSheetGenerator(root)
    root.mainloop()