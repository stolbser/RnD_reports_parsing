import os
import json
import pandas as pd
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import tkinter as tk
from tkinter import filedialog, messagebox


def add_image_group_2cols(doc, image_paths, img_width_cm=7):
    """
    Вставка группы изображений в таблицу 2 колонки.
    """
    if not image_paths:
        return None

    rows = (len(image_paths) + 1) // 2
    table = doc.add_table(rows=rows, cols=2)
    table.autofit = False

    idx = 0
    for r in range(rows):
        for c in range(2):
            if idx >= len(image_paths):
                cell = table.cell(r, c)
                cell._element.clear_content()
                continue

            cell = table.cell(r, c)
            cell._element.clear_content()
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            try:
                run = p.add_run()
                run.add_picture(image_paths[idx], width=Cm(img_width_cm))
            except Exception as e:
                p.add_run(f"[Ошибка загрузки: {os.path.basename(image_paths[idx])}]")
                print(f"⚠️ Ошибка загрузки изображения {image_paths[idx]}: {e}")

            idx += 1

    doc.add_paragraph()
    return table


def add_inclusion_groups_from_folders(doc, root_dir, fig_counter_start=1):
    """
    Вставка групп изображений из подпапок (по типам неметаллических включений).
    """
    fig_counter = fig_counter_start

    if not os.path.isdir(root_dir):
        p = doc.add_paragraph("[Папка с изображениями не найдена]")
        p.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return fig_counter

    folders = sorted([
        f for f in os.listdir(root_dir)
        if os.path.isdir(os.path.join(root_dir, f))
    ])

    if not folders:
        p = doc.add_paragraph("[Нет подпапок с изображениями]")
        p.italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return fig_counter

    for folder in folders:
        folder_path = os.path.join(root_dir, folder)

        images = sorted([
            os.path.join(folder_path, f)
            for f in os.listdir(folder_path)
            if f.lower().endswith(('.png', '.jpg', '.jpeg', '.tif', '.tiff', '.bmp'))
        ])

        if not images:
            continue

        add_image_group_2cols(doc, images)

        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f'Рисунок {fig_counter} – Распределение химических элементов в области неметаллического включения {folder}')
        run.bold = False

        fig_counter += 1
        doc.add_paragraph()

    return fig_counter


def make_report(data_dir):
    if not os.path.isdir(data_dir):
        raise FileNotFoundError(f"Директория с данными не найдена: {data_dir}")

    params_path = os.path.join(data_dir, "params.json.txt")
    if not os.path.isfile(params_path):
        raise FileNotFoundError(f"Файл параметров не найден: {params_path}")

    with open(params_path, encoding="utf-8") as f:
        params = json.load(f)

    script_dir = os.path.dirname(os.path.abspath(__file__))
    template_path = os.path.join(script_dir, "template.docx")

    if not os.path.isfile(template_path):
        template_path = "template.docx"
        if not os.path.isfile(template_path):
            raise FileNotFoundError(
                "Шаблон 'template.docx' не найден. Поместите его в папку со скриптом."
            )

    doc = Document(template_path)

    replacements = {
        "{{STEEL_GRADE}}": params.get("steel_grade", "—"),
        "{{HEAT_NUMBER}}": params.get("heat_number", "—"),
        "{{METHODS}}": params.get("methods_text", "")
    }

    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    fig_counter = 1
    fig_start = fig_counter

    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if "[FIGURES]" in paragraph.text:
            paragraphs_to_remove.append(paragraph)
            figures_dir = os.path.join(data_dir, "figures")
            if os.path.isdir(figures_dir):
                fig_counter = add_inclusion_groups_from_folders(doc, figures_dir, fig_counter)
            else:
                p = doc.add_paragraph("[Папка 'figures' не найдена]")
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            fig_end = fig_counter - 1 if fig_counter > 1 else 1

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    if not paragraphs_to_remove:
        fig_start = 1
        fig_end = 1

    for paragraph in doc.paragraphs:
        if "{{EDS_TEXT}}" in paragraph.text:
            text = params.get("eds_text", "")
            text = text.replace("{{FIG_START}}", str(fig_start))
            text = text.replace("{{FIG_END}}", str(fig_end))
            paragraph.text = paragraph.text.replace("{{EDS_TEXT}}", text)

        if "{{MORPH_TEXT}}" in paragraph.text:
            text = params.get("morphology_text", "")
            text = text.replace("{{FIG}}", str(fig_start))
            paragraph.text = paragraph.text.replace("{{MORPH_TEXT}}", text)

    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if "[TABLES]" in paragraph.text:
            paragraphs_to_remove.append(paragraph)

            tables_dir = os.path.join(data_dir, "tables")
            if not os.path.isdir(tables_dir):
                p = doc.add_paragraph("[Папка 'tables' не найдена]")
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            table_files = sorted([
                f for f in os.listdir(tables_dir)
                if f.lower().endswith(('.xlsx', '.xls'))
            ])

            if not table_files:
                p = doc.add_paragraph("[Нет таблиц для вставки]")
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                continue

            for idx, tbl_file in enumerate(table_files, 1):
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = caption.add_run(
                    f"Таблица {idx} – Количественный химический состав неметаллических включений"
                )
                run.bold = True

                try:
                    df = pd.read_excel(os.path.join(tables_dir, tbl_file))

                    table = doc.add_table(rows=1, cols=len(df.columns))
                    table.style = 'Light Grid Accent 1' if 'Light Grid Accent 1' in doc.styles else None

                    for j, column in enumerate(df.columns):
                        table.rows[0].cells[j].text = str(column)

                    for _, row in df.iterrows():
                        row_cells = table.add_row().cells
                        for j, value in enumerate(row):
                            row_cells[j].text = str(value) if pd.notna(value) else ""

                    doc.add_paragraph()

                except Exception as e:
                    p = doc.add_paragraph(f"[Ошибка загрузки таблицы '{tbl_file}': {str(e)}]")
                    p.italic = True
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    paragraphs_to_remove = []
    for i, paragraph in enumerate(doc.paragraphs):
        if "[CONCLUSIONS]" in paragraph.text:
            paragraphs_to_remove.append(paragraph)

            conclusions = params.get("conclusions", [])
            if conclusions:
                for item in conclusions:
                    doc.add_paragraph(f"• {item}")
            else:
                p = doc.add_paragraph("[Выводы отсутствуют]")
                p.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    output_dir = os.path.join(script_dir, "output")
    os.makedirs(output_dir, exist_ok=True)

    heat_num = params.get("heat_number", "unknown")
    output_path = os.path.join(output_dir, f"Заключение_{heat_num}.docx")

    doc.save(output_path)
    return output_path


def run_gui():
    def choose_folder():
        folder = filedialog.askdirectory(title="Выберите папку с данными плавки")
        if folder:
            entry.delete(0, tk.END)
            entry.insert(0, folder)

    def generate():
        data_path = entry.get().strip()
        if not data_path or not os.path.isdir(data_path):
            messagebox.showwarning("Внимание", "Выберите корректную папку с данными!")
            return

        try:
            output_file = make_report(data_path)
            messagebox.showinfo(
                "Готово",
                f"Отчёт успешно создан:\n{output_file}"
            )
            os.startfile(os.path.dirname(output_file))
        except Exception as e:
            messagebox.showerror("Ошибка", f"Не удалось создать отчёт:\n{str(e)}")
            import traceback
            traceback.print_exc()

    root = tk.Tk()
    root.title("Генератор отчётов НЛЗ")
    root.geometry("500x200")
    root.resizable(False, False)

    tk.Label(
        root,
        text="Папка с данными плавки:",
        font=("Arial", 10, "bold")
    ).pack(pady=(15, 5))

    entry = tk.Entry(root, width=60)
    entry.pack(padx=20)

    tk.Button(
        root,
        text="Выбрать папку",
        command=choose_folder,
        width=20
    ).pack(pady=10)

    tk.Button(
        root,
        text="Сформировать отчёт",
        command=generate,
        bg="#4CAF50",
        fg="white",
        width=25,
        height=2,
        font=("Arial", 10, "bold")
    ).pack(pady=10)

    tk.Label(
        root,
        text="Структура папки:\nparams.json.txt + папки figures/ и tables/",
        fg="gray",
        font=("Arial", 8)
    ).pack(side=tk.BOTTOM, pady=5)

    root.mainloop()


if __name__ == "__main__":
    run_gui()