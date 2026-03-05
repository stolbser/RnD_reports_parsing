import os
import glob
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn


# ============================================
# Функция для закрашивания ячеек
# ============================================
def shade_cell(cell, color='FFFF00'):
    """
    Закрашивает фон ячейки в указанный цвет
    color: HEX-код цвета (по умолчанию 'FFFF00' = жёлтый)
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Удаляем старую заливку, если есть
    shd = tcPr.xpath('./w:shd')
    if shd:
        tcPr.remove(shd[0])

    # Добавляем новую заливку
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


# ============================================
# Словарь для сокращения названий элементов
# ============================================
ELEMENT_SHORT_NAMES = {
    'Oxygen': 'O', 'Magnesium': 'Mg', 'Silicon': 'Si', 'Titanium': 'Ti', 'Manganese': 'Mn',
    'Iron': 'Fe', 'Aluminium': 'Al', 'Calcium': 'Ca', 'Sodium': 'Na', 'Potassium': 'K',
    'Chromium': 'Cr', 'Nickel': 'Ni', 'Copper': 'Cu', 'Zinc': 'Zn', 'Phosphorus': 'P',
    'Sulfur': 'S', 'Carbon': 'C', 'Nitrogen': 'N', 'Hydrogen': 'H', 'Chlorine': 'Cl',
    'Fluorine': 'F', 'Boron': 'B', 'Lithium': 'Li', 'Beryllium': 'Be', 'Vanadium': 'V',
    'Cobalt': 'Co', 'Arsenic': 'As', 'Strontium': 'Sr', 'Barium': 'Ba', 'Lead': 'Pb'
}


def get_short_name(full_name):
    """Преобразует полное название элемента в короткое (если известно), иначе возвращает оригинал"""
    return ELEMENT_SHORT_NAMES.get(full_name.strip(), full_name.strip())


def process_excel_file(filepath, doc):
    """
    Универсальная обработка Excel-файла с любым набором колонок
    """
    try:
        # Чтение Excel (без предположений о структуре)
        df = pd.read_excel(filepath)

        if df.empty:
            print(f"⚠️  Файл '{os.path.basename(filepath)}' пустой — пропущен")
            return False

        # Определение колонки со спектрами (первая колонка)
        spectrum_col = df.columns[0]
        data_cols = df.columns[1:]  # все остальные колонки — данные элементов

        if len(data_cols) == 0:
            print(f"⚠️  Файл '{os.path.basename(filepath)}' не содержит данных элементов — пропущен")
            return False

        # Удаление колонки 'Sum' (если существует)
        data_cols = [col for col in data_cols if str(col).strip().lower() != 'sum']
        df = df[[spectrum_col] + list(data_cols)].copy()

        # Удаление строк статистики (регистронезависимо)
        stats_keywords = ['mean', 'sigma', 'sigmamean', 'std', 'average', 'deviation']
        mask = df[spectrum_col].astype(str).str.lower().str.strip().isin(stats_keywords)
        df = df[~mask].copy()

        if df.empty:
            print(f"⚠️  Файл '{os.path.basename(filepath)}' содержит только статистику — пропущен")
            return False

        # Извлечение номера спектра (если возможно) или использование оригинального названия
        def extract_spectrum_number(val):
            import re
            val_str = str(val).strip()
            match = re.search(r'(\d+)', val_str)
            return match.group(1) if match else val_str

        df['№ спектра'] = df[spectrum_col].apply(extract_spectrum_number)

        # Подготовка колонок элементов: округление и форматирование
        short_names = {}
        for col in data_cols:
            short = get_short_name(str(col))
            short_names[col] = short
            # Округление и замена точки на запятую
            # Если значение < 0.5 → выводим прочерк "-"
            df[short] = pd.to_numeric(df[col], errors='coerce').round(1).apply(
                lambda x: '-' if pd.notna(x) and abs(x) < 0.5 else (
                    f"{x:.1f}".replace('.', ',') if pd.notna(x) else '-')
            )

        # === Расчёт суммы и определение строк для выделения ===
        numeric_cols = []
        for col in data_cols:
            short = get_short_name(str(col))
            numeric_cols.append(short)
            df[f'{short}_num'] = pd.to_numeric(df[col], errors='coerce').round(1)

        # Считаем сумму по строке (только числовые значения)
        df['row_sum'] = df[[f'{col}_num' for col in numeric_cols]].sum(axis=1).round(1)

        # Метка для выделения: сумма < 100
        df['highlight'] = (df['row_sum'] < 100) | (df['row_sum'] > 100)

        # Формирование итогового набора колонок
        element_cols = [short_names[col] for col in data_cols]
        display_cols = ['№ спектра'] + element_cols
        df_display = df[display_cols].copy()

        # === Добавление заголовка таблицы ===
        filename = os.path.splitext(os.path.basename(filepath))[0]
        doc.add_paragraph()
        heading = doc.add_heading(f"Таблица: {filename}", level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # === Создание основной таблицы ===
        num_cols = len(display_cols)
        table = doc.add_table(rows=1, cols=num_cols)
        table.style = 'Table Grid'

        # === Первая строка: объединённый заголовок элементов ===
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '№ спектра'
        hdr_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Объединение всех ячеек для элементов (кроме первой)
        if num_cols > 1:
            for i in range(1, num_cols):
                hdr_cells[i].text = ''
            hdr_cells[1].merge(hdr_cells[-1])
            hdr_cells[1].text = 'Содержание элементов, масс. %'
            hdr_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # === Вторая строка: названия элементов ===
        elem_cells = table.add_row().cells
        elem_cells[0].text = ''
        for i, elem in enumerate(element_cols):
            elem_cells[i + 1].text = elem
            elem_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # === Строки данных ===
        for idx, row in df_display.iterrows():
            row_cells = table.add_row().cells

            # Проверяем, нужно ли выделить эту строку
            highlight = df.loc[idx, 'highlight']

            # Первая ячейка: номер спектра
            row_cells[0].text = str(row['№ спектра'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if highlight:
                shade_cell(row_cells[0], 'FFFF00')  # жёлтый фон

            # Ячейки с элементами
            for i, elem in enumerate(element_cols):
                row_cells[i + 1].text = str(row[elem])
                row_cells[i + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                if highlight:
                    shade_cell(row_cells[i + 1], 'FFFF00')  # жёлтый фон

        # === Форматирование ячеек ===
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # === Добавление таблицы сумм ===
        doc.add_paragraph()
        sum_heading = doc.add_heading('Сумма по элементам', level=2)
        sum_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Создание таблицы сумм (2 колонки)
        sum_table = doc.add_table(rows=1, cols=2)
        sum_table.style = 'Table Grid'

        # Заголовок таблицы сумм
        sum_hdr = sum_table.rows[0].cells
        sum_hdr[0].text = '№ спектра'
        sum_hdr[1].text = 'Сумма, масс. %'

        # Форматирование заголовка
        for cell in sum_hdr:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True

        # Строки с суммами
        for idx, row in df.iterrows():
            row_cells = sum_table.add_row().cells
            highlight = row['highlight']

            # Номер спектра
            row_cells[0].text = str(row['№ спектра'])
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if highlight:
                shade_cell(row_cells[0], 'FFFF00')

            # Сумма (с заменой точки на запятую)
            sum_value = f"{row['row_sum']:.1f}".replace('.', ',')
            row_cells[1].text = sum_value
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if highlight:
                shade_cell(row_cells[1], 'FFFF00')

            # Форматирование ячеек
            for cell in row_cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)

        # Примечание о выделении
        highlight_count = df['highlight'].sum()
        if highlight_count > 0:
            doc.add_paragraph(f"* Жёлтым цветом выделены строки с суммой элементов < 100%")

        doc.add_paragraph()
        print(f"  ✅ Обработано: {len(df_display)} спектров, элементы: {', '.join(element_cols)}")
        if highlight_count > 0:
            print(f"     ⚠️  {highlight_count} строк(а) выделено жёлтым (сумма < 100%)")
        return True

    except Exception as e:
        print(f"  ❌ Ошибка при обработке '{os.path.basename(filepath)}': {e}")
        import traceback
        traceback.print_exc()
        return False


def process_folder(input_folder, output_file):
    """
    Обрабатывает все .xlsx файлы в папке и создаёт единый документ Word
    """
    # Поиск всех Excel-файлов в папке
    excel_files = sorted(glob.glob(os.path.join(input_folder, "*.xlsx")))

    if not excel_files:
        print(f"❌ В папке '{input_folder}' не найдено Excel-файлов (.xlsx)")
        return False

    print(f"📁 Найдено файлов: {len(excel_files)}")
    print("Список файлов:")
    for i, f in enumerate(excel_files, 1):
        print(f"  {i}. {os.path.basename(f)}")

    # Создание документа Word
    doc = Document()
    doc.add_heading('Анализ спектров', 0)
    doc.add_paragraph('Документ содержит таблицы, сформированные из всех Excel-файлов в папке.')
    doc.add_paragraph().add_run().add_break()

    # Обработка каждого файла
    success_count = 0
    for filepath in excel_files:
        print(f"\nОбработка: {os.path.basename(filepath)}")
        if process_excel_file(filepath, doc):
            success_count += 1

    # Сохранение результата
    doc.save(output_file)
    print(f"\n" + "=" * 60)
    print(f"✅ Готово! Успешно обработано {success_count} из {len(excel_files)} файлов.")
    print(f"📄 Документ сохранён: {os.path.abspath(output_file)}")
    print("=" * 60)
    return True


# ======================
# НАСТРОЙКА ПУТЕЙ ЗДЕСЬ ↓
# ======================
if __name__ == "__main__":
    # Папка с Excel-файлов (относительный или полный путь)
    INPUT_FOLDER = "C:/report_generator/tablegenerator/data/32g2fa"  # ← измените на имя вашей папки

    # Имя выходного файла Word
    OUTPUT_FILE = "C:/report_generator/tablegenerator/data/32g2fa/спектры_анализ.docx"

    # Запуск обработки
    process_folder(INPUT_FOLDER, OUTPUT_FILE)